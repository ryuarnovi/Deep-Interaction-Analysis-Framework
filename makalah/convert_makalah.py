#!/usr/bin/env python3
"""
Konversi Makalah_CEFR_Speech_Coach.md → .docx dan .pdf

Fitur:
  - Margin: Atas 4cm, Bawah 3cm, Kiri 4cm, Kanan 3cm
  - Font: Times New Roman 12pt, Spasi 1.5
  - Halaman Judul (Nama: Rizki Ardiansyah Novianto, NIM: A11.2024.15546)
  - Penomoran Halaman:
      * Halaman Judul: Tanpa nomor halaman.
      * Daftar Isi, Gambar, Tabel, Abstrak (Sebelum BAB 1): Romawi Kecil (i, ii, iii, iv, v, dst.).
      * BAB 1 s.d. Selesai: Angka Arab (1, 2, 3, dst.) dimulai kembali dari 1.
  - Hyperlink pada Daftar Isi, Daftar Gambar, Daftar Tabel.
  - Hyperlink pada Sitasi IEEE (contoh: [1] langsung melompat ke referensi [1] di Daftar Pustaka).
  - Format Bab Utama: "BAB 1\nPENDAHULUAN" (centered, bold, 14pt, uppercase).
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

BASE_DIR = "/Volumes/RYUARNOVI/project/ML/cefr-speech-coach/makalah"
INPUT_MD = os.path.join(BASE_DIR, "Makalah_CEFR_Speech_Coach.md")
OUTPUT_DOCX = os.path.join(BASE_DIR, "Makalah_CEFR_Speech_Coach.docx")
OUTPUT_PDF = os.path.join(BASE_DIR, "Makalah_CEFR_Speech_Coach.pdf")
IMAGES_DIR = "/Volumes/RYUARNOVI/project/ML/cefr-speech-coach/makalah/images"

FONT_BODY = 'Times New Roman'
FONT_SIZE = Pt(12)
LINE_SPACING = 1.5


# ====================================================================
# BOOKMARK & HYPERLINK HELPERS
# ====================================================================

_bookmark_counter = 0

def _next_bookmark_id():
    global _bookmark_counter
    _bookmark_counter += 1
    return _bookmark_counter

def add_bookmark(paragraph, bookmark_name):
    """Sisipkan bookmark (anchor) pada paragraph."""
    bid = str(_next_bookmark_id())
    tag_start = OxmlElement('w:bookmarkStart')
    tag_start.set(qn('w:id'), bid)
    tag_start.set(qn('w:name'), bookmark_name)
    paragraph._p.insert(0, tag_start)

    tag_end = OxmlElement('w:bookmarkEnd')
    tag_end.set(qn('w:id'), bid)
    paragraph._p.append(tag_end)

def add_hyperlink(paragraph, bookmark_name, text, bold=False, size=None, font=None, color_hex='1F4E79'):
    """Tambah hyperlink internal (klik → jump ke bookmark)."""
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    hyperlink.set(qn('w:history'), '1')

    run_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Font name
    rFonts = OxmlElement('w:rFonts')
    fn = font or FONT_BODY
    rFonts.set(qn('w:ascii'), fn)
    rFonts.set(qn('w:hAnsi'), fn)
    rFonts.set(qn('w:cs'), fn)
    rPr.append(rFonts)

    # Font size
    sz_val = size or FONT_SIZE
    sz_pt = int(sz_val.pt * 2)  # half-points
    rSz = OxmlElement('w:sz')
    rSz.set(qn('w:val'), str(sz_pt))
    rPr.append(rSz)
    rSzCs = OxmlElement('w:szCs')
    rSzCs.set(qn('w:val'), str(sz_pt))
    rPr.append(rSzCs)

    # Bold
    if bold:
        b_elem = OxmlElement('w:b')
        rPr.append(b_elem)

    # Color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), color_hex)
    rPr.append(color)

    # Underline
    u_elem = OxmlElement('w:u')
    u_elem.set(qn('w:val'), 'single')
    rPr.append(u_elem)

    run_elem.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.set(qn('xml:space'), 'preserve')
    text_elem.text = text
    run_elem.append(text_elem)

    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


# ====================================================================
# GENERAL HELPERS
# ====================================================================

def add_page_number_footer(section):
    """Nomor halaman di footer center."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    for field_type in ['begin', None, 'end']:
        run = p.add_run()
        run.font.size = Pt(11)
        run.font.name = FONT_BODY
        if field_type == 'begin':
            fc = OxmlElement('w:fldChar')
            fc.set(qn('w:fldCharType'), 'begin')
            run._r.append(fc)
        elif field_type is None:
            it = OxmlElement('w:instrText')
            it.set(qn('xml:space'), 'preserve')
            it.text = ' PAGE '
            run._r.append(it)
        else:
            fc = OxmlElement('w:fldChar')
            fc.set(qn('w:fldCharType'), 'end')
            run._r.append(fc)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def make_run(paragraph, text, bold=False, italic=False, size=None, font=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font or FONT_BODY
    run.font.size = size or FONT_SIZE
    if color:
        run.font.color.rgb = color
    return run


def _add_text_runs_with_citations(paragraph, text, bold=False, italic=False, size=None, font=None):
    """Split text by IEEE citation format like [1], [2] and add internal links to bibliography bookmarks."""
    tokens = re.split(r'(\[\d+\])', text)
    for token in tokens:
        if not token:
            continue
        match = re.match(r'^\[(\d+)\]$', token)
        if match:
            ref_num = match.group(1)
            # Hyperlink to bibliography bookmark ref_X
            add_hyperlink(paragraph, f'ref_{ref_num}', token, bold=bold, size=size, font=font, color_hex='0000D0')
        else:
            make_run(paragraph, token, bold=bold, italic=italic, size=size, font=font)


def add_formatted_text(paragraph, text, default_size=None, default_font=None):
    sz = default_size or FONT_SIZE
    fn = default_font or FONT_BODY
    parts = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*|`[^`]+?`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            _add_text_runs_with_citations(paragraph, part[2:-2], bold=True, size=sz, font=fn)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            _add_text_runs_with_citations(paragraph, part[1:-1], italic=True, size=sz, font=fn)
        elif part.startswith('`') and part.endswith('`'):
            r = make_run(paragraph, part[1:-1], size=Pt(9), font='Courier New')
            r.font.color.rgb = RGBColor(40, 40, 40)
        else:
            _add_text_runs_with_citations(paragraph, part, size=sz, font=fn)


def add_code_block(doc, code_lines):
    for code_line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Cm(1.27)
        r = p.add_run(code_line)
        r.font.name = 'Courier New'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(20, 20, 20)


def setup_page(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(4.0)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(3.0)


def add_body_paragraph(doc, text, first_indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = LINE_SPACING
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    add_formatted_text(p, text)
    return p


# ====================================================================
# HALAMAN JUDUL
# ====================================================================

def create_cover_page(doc):
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(p, 'MAKALAH', bold=True, size=Pt(16))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = LINE_SPACING
    make_run(p,
             'IMPLEMENTASI KECERDASAN BUATAN UNTUK PENILAIAN OTOMATIS\n'
             'KEMAMPUAN BERBICARA BAHASA INGGRIS BERBASIS KERANGKA CEFR:\n'
             'PENDEKATAN MACHINE LEARNING DAN DEEP LEARNING',
             bold=True, size=Pt(14))

    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = LINE_SPACING
    make_run(p, 'Disusun untuk memenuhi tugas\nLiterasi Informasi', size=Pt(12))

    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(p, 'Disusun oleh:', size=Pt(12))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    make_run(p, 'Rizki Ardiansyah Novianto', bold=True, size=Pt(14))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    make_run(p, 'NIM. A11.2024.15546', bold=True, size=Pt(12))

    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = LINE_SPACING
    make_run(p, 'PROGRAM STUDI TEKNIK INFORMATIKA\nFAKULTAS ILMU KOMPUTER\nUNIVERSITAS DIAN NUSWANTORO\n2026',
             bold=True, size=Pt(12))


# ====================================================================
# BAB MAPPING
# ====================================================================

CHAPTERS = [
    ('1 Pendahuluan',          1, 'bab1', 'PENDAHULUAN'),
    ('2 Tinjauan Pustaka',     2, 'bab2', 'TINJAUAN PUSTAKA'),
    ('3 Metodologi Penelitian',3, 'bab3', 'METODOLOGI PENELITIAN'),
    ('4 Hasil dan Pembahasan', 4, 'bab4', 'HASIL DAN PEMBAHASAN'),
    ('5 Kesimpulan',           5, 'bab5', 'KESIMPULAN'),
    ('Daftar Pustaka',         0, 'dafpus', 'DAFTAR PUSTAKA'),
]

def _bm_name(text):
    clean = re.sub(r'[^a-zA-Z0-9]', '_', text)[:40]
    return f'sec_{clean}'


# ====================================================================
# DAFTAR ISI / DAFTAR GAMBAR / DAFTAR TABEL
# ====================================================================

def _toc_entry(doc, text, bookmark, level=0, is_bab=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.left_indent = Cm(1.0) if level >= 1 else Cm(0)

    if bookmark:
        add_hyperlink(p, bookmark, text,
                      bold=(level == 0 or is_bab),
                      size=Pt(12), font=FONT_BODY)
    else:
        make_run(p, text, bold=(level == 0), size=Pt(12))


def create_daftar_isi(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    add_bookmark(p, 'daftar_isi')
    make_run(p, 'DAFTAR ISI', bold=True, size=Pt(14))

    _toc_entry(doc, 'DAFTAR ISI', 'daftar_isi', 0)
    _toc_entry(doc, 'DAFTAR GAMBAR', 'daftar_gambar', 0)
    _toc_entry(doc, 'DAFTAR TABEL', 'daftar_tabel', 0)
    _toc_entry(doc, 'ABSTRAK', 'abstrak', 0)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)

    _toc_entry(doc, 'BAB 1  PENDAHULUAN', 'bab1', 0, is_bab=True)
    _toc_entry(doc, '1.1  Latar Belakang', _bm_name('1.1 Latar Belakang'), 1)
    _toc_entry(doc, '1.2  Rumusan Masalah', _bm_name('1.2 Rumusan Masalah'), 1)
    _toc_entry(doc, '1.3  Tujuan Penelitian', _bm_name('1.3 Tujuan Penelitian'), 1)
    _toc_entry(doc, '1.4  Manfaat Penelitian', _bm_name('1.4 Manfaat Penelitian'), 1)

    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)

    _toc_entry(doc, 'BAB 2  TINJAUAN PUSTAKA', 'bab2', 0, is_bab=True)
    _toc_entry(doc, '2.1  Kecerdasan Buatan (Artificial Intelligence)', _bm_name('2.1 Kecerdasan Buatan'), 1)
    _toc_entry(doc, '2.2  Machine Learning', _bm_name('2.2 Machine Learning'), 1)
    _toc_entry(doc, '2.3  Deep Learning', _bm_name('2.3 Deep Learning'), 1)
    _toc_entry(doc, '2.4  Penilaian Kemampuan Berbicara', _bm_name('2.4 Penilaian Kemampuan Berbicara'), 1)
    _toc_entry(doc, '2.5  Kerangka CEFR', _bm_name('2.5 Kerangka CEFR'), 1)
    _toc_entry(doc, '2.6  Penilaian Pengucapan', _bm_name('2.6 Penilaian Pengucapan'), 1)
    _toc_entry(doc, '2.7  Penilaian Otomatis Kemampuan Berbicara', _bm_name('2.7 Penilaian Otomatis Kemampuan Berbicara'), 1)

    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)

    _toc_entry(doc, 'BAB 3  METODOLOGI PENELITIAN', 'bab3', 0, is_bab=True)
    _toc_entry(doc, '3.1  Machine Learning (Part A)', _bm_name('3.1 Machine Learning'), 1)
    _toc_entry(doc, '3.2  Deep Learning (Part B)', _bm_name('3.2 Deep Learning'), 1)

    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)

    _toc_entry(doc, 'BAB 4  HASIL DAN PEMBAHASAN', 'bab4', 0, is_bab=True)
    _toc_entry(doc, '4.1  Eksplorasi Data (EDA)', _bm_name('4.1 Eksplorasi Data'), 1)
    _toc_entry(doc, '4.2  Hasil Klasifikasi Machine Learning', _bm_name('4.2 Hasil Klasifikasi Machine Learning'), 1)
    _toc_entry(doc, '4.3  Hasil Regresi Skor Prosodi', _bm_name('4.3 Hasil Regresi Skor Prosodi'), 1)
    _toc_entry(doc, '4.4  Hasil Deep Learning — Wav2Vec2', _bm_name('4.4 Hasil Deep Learning'), 1)
    _toc_entry(doc, '4.5  Perbandingan ML vs. Deep Learning', _bm_name('4.5 Perbandingan ML vs'), 1)
    _toc_entry(doc, '4.6  Daftar Library yang Digunakan', _bm_name('4.6 Daftar Library'), 1)
    _toc_entry(doc, '4.7  Pipeline Lengkap Sistem', _bm_name('4.7 Pipeline Lengkap Sistem'), 1)

    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)

    _toc_entry(doc, 'BAB 5  KESIMPULAN', 'bab5', 0, is_bab=True)

    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)

    _toc_entry(doc, 'DAFTAR PUSTAKA', 'dafpus', 0, is_bab=True)


FIGURES = [
    ('Gambar 4.1', 'Distribusi CEFR Level dan Skor Prosody Similarity'),
    ('Gambar 4.2', 'Confusion Matrix Deep MLP dan Perbandingan Macro F1'),
    ('Gambar 4.3', 'True vs. Predicted Score dan Top 10 Feature Importance'),
    ('Gambar 4.4', 'Confusion Matrix Wav2Vec2 dan F1 Score per Epoch'),
    ('Gambar 4.5', 'Perbandingan F1 Score: ML Traditional vs. Deep Learning'),
]

def create_daftar_gambar(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    add_bookmark(p, 'daftar_gambar')
    make_run(p, 'DAFTAR GAMBAR', bold=True, size=Pt(14))

    for fig_id, fig_title in FIGURES:
        bm = fig_id.replace(' ', '_').replace('.', '_')
        full_text = f'{fig_id}  {fig_title}'
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_hyperlink(p, bm, full_text, size=Pt(12), font=FONT_BODY)


TABLES_LIST = [
    ('Tabel 3.1', 'Daftar Fitur Dataset Prosodi'),
    ('Tabel 3.2', 'Konfigurasi Model Klasifikasi ML'),
    ('Tabel 3.3', 'Distribusi Kelas Data Audio'),
    ('Tabel 3.4', 'Konfigurasi LoRA'),
    ('Tabel 3.5', 'Bobot Kelas (Class Weights)'),
    ('Tabel 3.6', 'Hyperparameter Training Deep Learning'),
    ('Tabel 4.1', 'Distribusi Tingkat CEFR pada Dataset Prosodi'),
    ('Tabel 4.2', 'Perbandingan Performa Model Klasifikasi ML'),
    ('Tabel 4.3', 'Classification Report Model Deep MLP'),
    ('Tabel 4.4', 'Metrik Evaluasi Model Regresi'),
    ('Tabel 4.5', 'Top 10 Feature Importance (Regresi)'),
    ('Tabel 4.6', 'Riwayat Training Wav2Vec2 per Epoch'),
    ('Tabel 4.7', 'Hasil Evaluasi Akhir Wav2Vec2'),
    ('Tabel 4.8', 'Classification Report Model Wav2Vec2'),
    ('Tabel 4.9', 'Ringkasan Perbandingan ML vs. DL'),
    ('Tabel 4.10', 'Library dan Dependensi'),
]

def create_daftar_tabel(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    add_bookmark(p, 'daftar_tabel')
    make_run(p, 'DAFTAR TABEL', bold=True, size=Pt(14))

    for tbl_id, tbl_title in TABLES_LIST:
        bm = tbl_id.replace(' ', '_').replace('.', '_')
        full_text = f'{tbl_id}  {tbl_title}'
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_hyperlink(p, bm, full_text, size=Pt(12), font=FONT_BODY)


# ====================================================================
# HEADING HELPERS
# ====================================================================

def add_heading_bab(doc, bab_number, display_title, bookmark_name):
    if bab_number > 0:
        p1 = doc.add_paragraph(style='Heading 1')
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(24)
        p1.paragraph_format.space_after = Pt(6)
        p1.paragraph_format.line_spacing = LINE_SPACING
        add_bookmark(p1, bookmark_name)
        make_run(p1, f'BAB {bab_number}', bold=True, size=Pt(14))

        p2 = doc.add_paragraph(style='Heading 1')
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(18)
        p2.paragraph_format.line_spacing = LINE_SPACING
        make_run(p2, display_title, bold=True, size=Pt(14))
    else:
        p1 = doc.add_paragraph(style='Heading 1')
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(24)
        p1.paragraph_format.space_after = Pt(18)
        p1.paragraph_format.line_spacing = LINE_SPACING
        add_bookmark(p1, bookmark_name)
        make_run(p1, display_title, bold=True, size=Pt(14))


def add_heading_sub(doc, text):
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.keep_with_next = True
    bm = _bm_name(text)
    add_bookmark(p, bm)
    make_run(p, text, bold=True, size=Pt(12))
    return p


def add_heading_subsub(doc, text):
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.keep_with_next = True
    make_run(p, text, bold=True, size=Pt(12))
    return p


# ====================================================================
# MAIN CONVERTER
# ====================================================================

def convert_md_to_docx(input_md, output_docx):
    if not os.path.exists(input_md):
        print(f"❌ File {input_md} tidak ditemukan.")
        return False

    doc = Document()

    # ================================================================
    # SECTION 1: HALAMAN JUDUL (Cover - Tanpa Nomor Halaman)
    # ================================================================
    section1 = doc.sections[0]
    setup_page(section1)
    # Jadikan footer kosong
    section1.footer.is_linked_to_previous = False
    p_empty = section1.footer.paragraphs[0] if section1.footer.paragraphs else section1.footer.add_paragraph()
    p_empty.text = ""

    create_cover_page(doc)

    # ================================================================
    # SECTION 2: DAFTAR ISI, GAMBAR, TABEL, ABSTRAK (Romawi ii, iii, dst.)
    # ================================================================
    section2 = doc.add_section(WD_SECTION_START.NEW_PAGE)
    setup_page(section2)
    section2.footer.is_linked_to_previous = False

    # Atur tipe penomoran Romawi Kecil
    sectPr2 = section2._sectPr
    pgNumType2 = OxmlElement('w:pgNumType')
    pgNumType2.set(qn('w:fmt'), 'lowerRoman')
    sectPr2.append(pgNumType2)

    add_page_number_footer(section2)

    create_daftar_isi(doc)
    add_page_break(doc)

    create_daftar_gambar(doc)
    add_page_break(doc)

    create_daftar_tabel(doc)
    add_page_break(doc)

    # ================================================================
    # PREPARE READING MARKDOWN CONTENT
    # ================================================================

    gambar_map = {
        'Gambar 4.1': os.path.join(IMAGES_DIR, 'eda_distributions.png'),
        'Gambar 4.2': os.path.join(IMAGES_DIR, 'partA_results.png'),
        'Gambar 4.3': os.path.join(IMAGES_DIR, 'partA_regression.png'),
        'Gambar 4.4': os.path.join(IMAGES_DIR, 'partB_results.png'),
        'Gambar 4.5': os.path.join(IMAGES_DIR, 'final_comparison.png'),
    }

    chapter_lookup = {}
    for md_text, bab_num, bm, title in CHAPTERS:
        chapter_lookup[md_text] = (bab_num, bm, title)

    with open(input_md, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    bab_count = 0
    in_references = False
    pending_table_caption = None

    # Hapus heading styles default color dan atur font akademik
    for hlvl in range(1, 5):
        hn = f'Heading {hlvl}'
        if hn in doc.styles:
            style_h = doc.styles[hn]
            style_h.font.name = FONT_BODY
            style_h.font.color.rgb = RGBColor(0, 0, 0)
            style_h.font.bold = True
            if hlvl == 1:
                style_h.font.size = Pt(14)
            else:
                style_h.font.size = Pt(12)

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == '---':
            i += 1
            continue

        # ---- Code blocks ----
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(3)
            add_code_block(doc, code_lines)
            sp2 = doc.add_paragraph()
            sp2.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # ---- Math blocks ----
        if stripped.startswith('$$') and stripped.endswith('$$'):
            formula = stripped.strip('$ ').strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(formula)
            r.italic = True
            r.font.size = Pt(12)
            r.font.name = 'Cambria Math'
            i += 1
            continue

        # ---- Tables ----
        if stripped.startswith('|') and i + 1 < len(lines):
            next_l = lines[i+1].strip() if i+1 < len(lines) else ''
            if '|---' in next_l or '|:---' in next_l or '| :---' in next_l:
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    rt = lines[i].strip()
                    if not re.match(r'^\|[\s:\-|]+\|$', rt):
                        table_lines.append(rt)
                    i += 1

                if table_lines:
                    rows_data = []
                    for tl in table_lines:
                        cells = [c.strip() for c in tl.split('|')]
                        if cells and cells[0] == '': cells = cells[1:]
                        if cells and cells[-1] == '': cells = cells[:-1]
                        rows_data.append(cells)

                    if rows_data:
                        num_cols = max(len(r) for r in rows_data)
                        table = doc.add_table(rows=0, cols=num_cols)
                        table.style = 'Table Grid'
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        table.autofit = True

                        for r_idx, row_cells in enumerate(rows_data):
                            row = table.add_row()
                            for c_idx, cv in enumerate(row_cells):
                                if c_idx < len(row.cells):
                                    cell = row.cells[c_idx]
                                    p = cell.paragraphs[0]
                                    p.paragraph_format.space_before = Pt(2)
                                    p.paragraph_format.space_after = Pt(2)
                                    p.paragraph_format.line_spacing = 1.0
                                    add_formatted_text(p, cv,
                                                       default_size=Pt(10),
                                                       default_font=FONT_BODY)
                                    if r_idx == 0:
                                        for run in p.runs:
                                            run.bold = True
                                        shading = parse_xml(
                                            f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
                                        cell._tc.get_or_add_tcPr().append(shading)

                        # Render pending table caption below the table
                        if pending_table_caption:
                            tbl_id, tbl_title, bm = pending_table_caption
                            p_cap = doc.add_paragraph()
                            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_cap.paragraph_format.space_before = Pt(6)
                            p_cap.paragraph_format.space_after = Pt(12)
                            add_bookmark(p_cap, bm)
                            make_run(p_cap, f'{tbl_id} — {tbl_title}', bold=True, size=Pt(10))
                            pending_table_caption = None

                        sp = doc.add_paragraph()
                        sp.paragraph_format.space_after = Pt(6)
                continue

        # ---- Image embeds ![alt](path) ----
        img_match = re.match(r'^!\[(.+?)\]\((.+?)\)$', stripped)
        if img_match:
            # Skip this line as it will be processed and rendered by the Caption handler
            i += 1
            continue

        # ================================================================
        # HEADINGS
        # ================================================================

        # H1: Judul makalah (Abstrak area)
        if stripped.startswith('# ') and not stripped.startswith('## '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = LINE_SPACING
            add_bookmark(p, 'abstrak')
            make_run(p, stripped[2:], bold=True, size=Pt(13))
            i += 1
            continue

        # H2: Bab utama (## 1 Pendahuluan, ## 2 Tinjauan Pustaka, dst.)
        if stripped.startswith('## ') and not stripped.startswith('### '):
            heading_text = stripped[3:]
            ch_info = chapter_lookup.get(heading_text)

            if ch_info:
                bab_num, bm, title = ch_info
                bab_count += 1

                # ============================================================
                # SECTION 3: BAB 1 s.d. Selesai (Angka Arab dimulai kembali dari 1)
                # ============================================================
                if bab_count == 1:
                    # Buat Section 3 saat memasuki BAB 1
                    section3 = doc.add_section(WD_SECTION_START.NEW_PAGE)
                    setup_page(section3)
                    section3.footer.is_linked_to_previous = False

                    # Atur tipe penomoran Angka Arab (decimal) dimulai dari 1
                    sectPr3 = section3._sectPr
                    pgNumType3 = OxmlElement('w:pgNumType')
                    pgNumType3.set(qn('w:fmt'), 'decimal')
                    pgNumType3.set(qn('w:start'), '1')
                    sectPr3.append(pgNumType3)

                    add_page_number_footer(section3)
                else:
                    # Bab-bab berikutnya cukup page break normal dalam Section 3
                    add_page_break(doc)

                # Set flag in_references jika bab Daftar Pustaka
                if title == 'DAFTAR PUSTAKA':
                    in_references = True

                add_heading_bab(doc, bab_num, title, bm)
            else:
                # Non-chapter H2
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(12)
                make_run(p, heading_text.upper(), bold=True, size=Pt(14))

            i += 1
            continue

        # H3: Sub-bab
        if stripped.startswith('### ') and not stripped.startswith('#### '):
            add_heading_sub(doc, stripped[4:])
            i += 1
            continue

        # H4: Sub-sub-bab
        if stripped.startswith('#### '):
            add_heading_subsub(doc, stripped[5:])
            i += 1
            continue

        # ---- Gambar caption ----
        gambar_key = None
        for gk in gambar_map:
            if stripped.startswith(f'**{gk}**'):
                gambar_key = gk
                break

        if gambar_key:
            img_path = gambar_map[gambar_key]
            caption_text = stripped.replace('**', '')
            bm = gambar_key.replace(' ', '_').replace('.', '_')

            # Render Image first
            if os.path.exists(img_path):
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.paragraph_format.space_before = Pt(12)
                img_p.paragraph_format.space_after = Pt(4)
                run = img_p.add_run()
                run.add_picture(img_path, width=Cm(13.0))

            # Render Caption below the image
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.paragraph_format.space_before = Pt(4)
            cap_p.paragraph_format.space_after = Pt(12)
            add_bookmark(cap_p, bm)
            make_run(cap_p, caption_text, bold=True, size=Pt(10))

            i += 1
            continue

        # ---- Tabel caption "**Tabel X.Y — ...**" ----
        tabel_match = re.match(r'^\*\*(Tabel \d+\.\d+)\s*[—–-]\s*(.+?)\*\*$', stripped)
        if tabel_match:
            tbl_id = tabel_match.group(1)
            tbl_title = tabel_match.group(2)
            bm = tbl_id.replace(' ', '_').replace('.', '_')
            pending_table_caption = (tbl_id, tbl_title, bm)
            i += 1
            continue

        # ---- IEEE Bibliography Entry in References (e.g. [1] M. Amrate...) ----
        ref_p_match = re.match(r'^\[(\d+)\]\s*(.+)$', stripped)
        if ref_p_match and in_references:
            ref_num = ref_p_match.group(1)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = LINE_SPACING
            # Add citation target bookmark ref_X
            add_bookmark(p, f'ref_{ref_num}')
            # Format [X] bold and the rest standard
            make_run(p, f'[{ref_num}] ', bold=True)
            add_formatted_text(p, ref_p_match.group(2))
            i += 1
            continue

        # ---- Other bold-only lines ----
        if stripped.startswith('**') and stripped.endswith('**'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            make_run(p, stripped[2:-2], bold=True, size=Pt(11))
            i += 1
            continue

        # ---- Bullet lists ----
        if stripped.startswith('- ') or (stripped.startswith('* ') and not stripped.startswith('**')):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = LINE_SPACING
            add_formatted_text(p, stripped[2:])
            i += 1
            continue

        # ---- Numbered lists ----
        num_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if num_match:
            num_val = num_match.group(1)
            list_content = num_match.group(2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = LINE_SPACING
            # Hanging indent formatting:
            p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.first_line_indent = Cm(-0.75)

            make_run(p, f"{num_val}.  ", bold=True)
            add_formatted_text(p, list_content)
            i += 1
            continue

        # ---- Normal paragraph ----
        needs_indent = not stripped.startswith('[') and not stripped.startswith('di mana')
        add_body_paragraph(doc, stripped, first_indent=needs_indent)
        i += 1

    # ---- Save ----
    doc.save(output_docx)
    sz = os.path.getsize(output_docx) / 1024
    print(f"✅ DOCX berhasil dibuat: {output_docx} ({sz:.0f} KB)")
    return True


# ====================================================================
# PDF
# ====================================================================

def convert_docx_to_pdf(docx_path, pdf_path):
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            sz = os.path.getsize(pdf_path) / 1024
            print(f"✅ PDF berhasil dibuat: {pdf_path} ({sz:.0f} KB)")
            return True
    except Exception as e:
        print(f"⚠️  docx2pdf error: {e}")

    import subprocess
    try:
        abs_docx = os.path.abspath(docx_path)
        abs_pdf = os.path.abspath(pdf_path)
        script = f'''
        tell application "Microsoft Word"
            activate
            open "{abs_docx}"
            set theDoc to active document
            save as theDoc file name "{abs_pdf}" file format format PDF
            close theDoc saving no
        end tell
        '''
        subprocess.run(['osascript', '-e', script],
                        capture_output=True, text=True, timeout=120)
        if os.path.exists(pdf_path):
            sz = os.path.getsize(pdf_path) / 1024
            print(f"✅ PDF berhasil dibuat (AppleScript): {pdf_path} ({sz:.0f} KB)")
            return True
    except Exception as e:
        print(f"⚠️  AppleScript error: {e}")

    print("⚠️  Silakan buka DOCX lalu Save As → PDF secara manual.")
    print(f"   {docx_path}")
    return False


# ====================================================================
# MAIN
# ====================================================================

if __name__ == '__main__':
    print("=" * 60)
    print("  Konversi Makalah CEFR Speech Coach → DOCX + PDF")
    print("  Format: Margin 4-3-4-3, TNR 12pt, Spasi 1.5")
    print("  Fitur: Cover, Daftar Isi (Romawi), BAB I (Arab, Mulai 1), IEEE Link")
    print("=" * 60)

    success = convert_md_to_docx(INPUT_MD, OUTPUT_DOCX)
    if success:
        print("\n📄 Konversi ke PDF...")
        convert_docx_to_pdf(OUTPUT_DOCX, OUTPUT_PDF)

    print("\n✅ Selesai!")
